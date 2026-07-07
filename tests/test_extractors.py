import pytest
import os
import tempfile
from app.extractors.factory import get_extractor, extract_from_file
from app.extractors.txt import TXTExtractor
from app.extractors.docx import DOCXExtractor
from app.extractors.pdf import PDFExtractor


class TestExtractorFactory:

    def test_returns_pdf_extractor_for_pdf(self):
        extractor = get_extractor("resume.pdf")
        assert isinstance(extractor, PDFExtractor)

    def test_returns_docx_extractor_for_docx(self):
        extractor = get_extractor("resume.docx")
        assert isinstance(extractor, DOCXExtractor)

    def test_returns_txt_extractor_for_txt(self):
        extractor = get_extractor("resume.txt")
        assert isinstance(extractor, TXTExtractor)

    def test_case_insensitive_extension(self):
        extractor = get_extractor("resume.PDF")
        assert isinstance(extractor, PDFExtractor)

    def test_unsupported_extension_raises_value_error(self):
        with pytest.raises(ValueError) as exc_info:
            get_extractor("resume.png")
        assert "Unsupported file type" in str(exc_info.value)

    def test_extract_from_file_handles_unsupported_gracefully(self):
        result = extract_from_file("resume.png")
        assert result["success"] is False
        assert "Unsupported file type" in result["error"]


class TestTXTExtractor:

    def test_extracts_text_from_valid_file(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as f:
            f.write("John Doe\nPython Developer\njohn@example.com")
            path = f.name

        try:
            extractor = TXTExtractor()
            result = extractor.extract(path)
            assert result["success"] is True
            assert "John Doe" in result["text"]
            assert result["error"] is None
        finally:
            os.unlink(path)

    def test_returns_error_for_missing_file(self):
        extractor = TXTExtractor()
        result = extractor.extract("nonexistent_file.txt")
        assert result["success"] is False
        assert result["error"] is not None

    def test_returns_error_for_empty_file(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False
        ) as f:
            f.write("")
            path = f.name

        try:
            extractor = TXTExtractor()
            result = extractor.extract(path)
            assert result["success"] is False
            assert "empty" in result["error"].lower()
        finally:
            os.unlink(path)

    def test_result_has_required_keys(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False
        ) as f:
            f.write("some content")
            path = f.name

        try:
            extractor = TXTExtractor()
            result = extractor.extract(path)
            assert "success" in result
            assert "text" in result
            assert "page_count" in result
            assert "error" in result
        finally:
            os.unlink(path)


class TestDOCXExtractor:

    def test_returns_error_for_missing_file(self):
        extractor = DOCXExtractor()
        result = extractor.extract("nonexistent.docx")
        assert result["success"] is False
        assert result["error"] is not None

    def test_extracts_text_from_valid_docx(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Python Developer")
        doc.add_paragraph("john@example.com")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name

        try:
            doc.save(path)
            extractor = DOCXExtractor()
            result = extractor.extract(path)
            assert result["success"] is True
            assert "John Doe" in result["text"]
            assert result["error"] is None
        finally:
            os.unlink(path)


class TestPDFExtractor:

    def test_returns_error_for_missing_file(self):
        extractor = PDFExtractor()
        result = extractor.extract("nonexistent.pdf")
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_extracts_text_from_real_pdf(self):
        path = "uploads/Mohan_D_Resume_updated.pdf"
        if not os.path.exists(path):
            pytest.skip("Real PDF not available for testing")

        extractor = PDFExtractor()
        result = extractor.extract(path)
        assert result["success"] is True
        assert len(result["text"]) > 100
        assert result["page_count"] >= 1