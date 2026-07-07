import os
from docx import Document
from app.extractors.base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """
    Extracts text from .docx files using python-docx.

    DOCX files are structured XML under the hood — python-docx
    reads the XML and gives us paragraphs and tables directly.
    This is more reliable than PDF extraction because DOCX
    preserves semantic structure (paragraphs, headings, tables).
    """

    def extract(self, file_path: str) -> dict:
        if not os.path.exists(file_path):
            return {
                "success": False,
                "text": "",
                "page_count": 0,
                "error": f"File not found: {file_path}"
            }

        try:
            doc = Document(file_path)

            extracted_lines = []

            # Extract text from paragraphs
            # Each paragraph is a block of text — headings, body text, bullets
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    extracted_lines.append(text)

            # Extract text from tables
            # Tables in DOCX are separate from paragraphs
            # We read cell by cell, row by row
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        # Join cells with pipe separator
                        extracted_lines.append(" | ".join(row_text))

            full_text = "\n".join(extracted_lines)

            if not full_text.strip():
                return {
                    "success": False,
                    "text": "",
                    "page_count": 1,
                    "error": "No text could be extracted from this DOCX file."
                }

            return {
                "success": True,
                "text": full_text,
                "page_count": 1,
                "error": None
            }

        except Exception as e:
            return {
                "success": False,
                "text": "",
                "page_count": 0,
                "error": f"Failed to extract text from DOCX: {str(e)}"
            }